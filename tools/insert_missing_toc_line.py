from __future__ import annotations

import argparse
import copy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def make_toc_paragraph(style_id: str, title: str, page: str):
    p = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    p_style = OxmlElement("w:pStyle")
    p_style.set(qn("w:val"), style_id)
    p_pr.append(p_style)
    p.append(p_pr)

    r_title = OxmlElement("w:r")
    t_title = OxmlElement("w:t")
    t_title.text = title
    r_title.append(t_title)
    p.append(r_title)

    r_tab = OxmlElement("w:r")
    r_tab.append(OxmlElement("w:tab"))
    p.append(r_tab)

    r_page = OxmlElement("w:r")
    t_page = OxmlElement("w:t")
    t_page.text = page
    r_page.append(t_page)
    p.append(r_page)

    return p


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    args = parser.parse_args()

    doc = Document(args.docx)
    existing = [
        " ".join(p.text.split())
        for p in doc.paragraphs
        if p.style.name.startswith("toc")
    ]
    if any(text.startswith("3.2.1. <Vehicle VIP Check-in>") for text in existing):
        print("toc_line_inserted=0")
        return

    inserted = 0
    for paragraph in doc.paragraphs:
        text = " ".join(paragraph.text.split())
        if paragraph.style.name.startswith("toc") and text.startswith("3.2. Use Case Specifications"):
            style_id = paragraph.style.style_id
            new_p = make_toc_paragraph(style_id, "3.2.1. <Vehicle VIP Check-in>", "17")
            paragraph._p.addnext(new_p)
            inserted = 1
            break

    doc.save(args.docx)
    print(f"toc_line_inserted={inserted}")


if __name__ == "__main__":
    main()
